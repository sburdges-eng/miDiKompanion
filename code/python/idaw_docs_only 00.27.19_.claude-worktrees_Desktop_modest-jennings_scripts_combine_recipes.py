#!/usr/bin/env python3
"""
Combine Lariat Recipe Books
Merges multiple recipe Word documents, detecting duplicates and versioning them
"""

import os
import hashlib
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
from datetime import datetime
import re

# Recipe source files - prioritize the 2025 FINAL version
RECIPE_FILES = [
    # Primary source - well formatted with proper headings
    "/Users/seanburdges/Library/CloudStorage/GoogleDrive-sburdges@gmail.com/My Drive/Lariat recipe 2025 FINAL.docx",
    # Additional sources (may have older/different versions)
    "/Users/seanburdges/Desktop/LARIAT-BIBLE-OFFICIAL-V.01.00/LARIAT/Lariat Recipe Book.docx",
]

# Only use alternative extraction if the primary source has issues
USE_ALTERNATIVE_EXTRACTION = False

OUTPUT_FILE = "/Users/seanburdges/.claude-worktrees/Desktop/modest-jennings/data/Lariat_Recipe_Book_Combined.docx"


def extract_recipes_from_doc(filepath):
    """Extract individual recipes from a Word document"""
    recipes = []
    
    if not os.path.exists(filepath):
        print(f"  ⚠️  File not found: {filepath}")
        return recipes
    
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  ❌ Error reading {filepath}: {e}")
        return recipes
    
    source = os.path.basename(filepath)
    
    current_recipe = {
        'title': None,
        'content': [],
        'source': source
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
        
        # Check if this is a heading (recipe title)
        is_heading = False
        
        if para.style and 'Heading' in para.style.name:
            is_heading = True
        
        # Skip the main "Lariat Recipe Book" title
        if 'lariat recipe book' in text.lower():
            continue
        
        if is_heading and len(text) > 3:
            # Save previous recipe if exists
            if current_recipe['title'] and current_recipe['content']:
                recipes.append(current_recipe.copy())
            
            # Start new recipe
            current_recipe = {
                'title': text.strip(),
                'content': [],
                'source': source
            }
        else:
            if current_recipe['title']:
                current_recipe['content'].append(text)
    
    # Don't forget the last recipe
    if current_recipe['title'] and current_recipe['content']:
        recipes.append(current_recipe)
    
    # Only try alternative extraction if enabled and few recipes found
    if USE_ALTERNATIVE_EXTRACTION and len(recipes) < 5:
        alt_recipes = extract_recipes_alternative(doc, source)
        if len(alt_recipes) > len(recipes):
            recipes = alt_recipes
    
    return recipes


def extract_recipes_alternative(doc, source):
    """Alternative extraction for documents without proper headings"""
    recipes = []
    
    current_recipe = {
        'title': None,
        'content': [],
        'source': source
    }
    
    # Recipe name keywords - must be in title to be considered a recipe
    recipe_keywords = ['sauce', 'brine', 'flour', 'rub', 'queso', 'cheese', 'batter',
                       'dressing', 'marinade', 'confit', 'jus', 'butter', 'spread',
                       'aioli', 'soup', 'chili', 'bread', 'melt', 'slaw', 'toast',
                       'succotash', 'salsa', 'glaze', 'oil', 'pickle', 'pico', 'relish',
                       'vinaigrette', 'seasoning', 'pepitas', 'jam', 'mustard']
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
            
        text_lower = text.lower()
        
        # Skip headers and common subsections
        if 'lariat recipe book' in text_lower:
            continue
        if text_lower in ['ingredients', 'directions', 'procedure', 'yield', 'method']:
            current_recipe['content'].append(text)
            continue
        
        # Check if all runs are bold (indicates title)
        is_bold_title = para.runs and all(run.bold for run in para.runs if run.text.strip())
        
        # Check for recipe keywords and proper format
        has_keyword = any(kw in text_lower for kw in recipe_keywords)
        proper_length = 4 < len(text) < 50  # Tighter length limit
        starts_capital = text[0].isupper() if text else False
        
        # More strict instruction filtering
        instruction_starts = ['add', 'mix', 'stir', 'cook', 'pour', 'combine', 'whisk', 
                              'blend', 'in a', 'on a', 'for a', 'for the', 'let', 'place',
                              'heat', 'bring', 'simmer', 'reduce', 'season', 'adjust']
        not_instruction = not any(text_lower.startswith(s) for s in instruction_starts)
        not_quantity = not re.match(r'^[\d½¼¾⅓⅔⅛]+', text)
        
        # Title shouldn't contain sentences (periods in middle)
        not_sentence = '.' not in text[:-1] if text else True
        
        is_recipe_title = (is_bold_title and has_keyword and proper_length and 
                          starts_capital and not_instruction and not_quantity and not_sentence)
        
        if is_recipe_title:
            # Save previous recipe
            if current_recipe['title'] and current_recipe['content']:
                recipes.append(current_recipe.copy())
            
            current_recipe = {
                'title': text.strip(),
                'content': [],
                'source': source
            }
        else:
            if current_recipe['title']:
                current_recipe['content'].append(text)
    
    if current_recipe['title'] and current_recipe['content']:
        recipes.append(current_recipe)
    
    return recipes


def normalize_title(title):
    """Normalize recipe title for comparison"""
    normalized = title.lower().strip()
    # Remove special characters but keep important ones
    normalized = re.sub(r'[^\w\s\-–]', '', normalized)
    normalized = re.sub(r'\s+', ' ', normalized)
    return normalized


def get_content_hash(content_list):
    """Create a hash of recipe content for duplicate detection"""
    content_str = '\n'.join(content_list).lower()
    content_str = re.sub(r'\s+', ' ', content_str)
    return hashlib.md5(content_str.encode()).hexdigest()[:12]


def combine_recipes():
    """Main function to combine all recipe books"""
    print("=" * 60)
    print("🍳 LARIAT RECIPE BOOK COMBINER")
    print("=" * 60)
    print(f"\nProcessing {len(RECIPE_FILES)} source files...\n")
    
    all_recipes = []
    
    # Extract recipes from each file
    for filepath in RECIPE_FILES:
        if not os.path.exists(filepath):
            print(f"⚠️  Skipping (not found): {os.path.basename(filepath)}")
            continue
        print(f"📖 Reading: {os.path.basename(filepath)}")
        recipes = extract_recipes_from_doc(filepath)
        print(f"   Found {len(recipes)} recipes")
        all_recipes.extend(recipes)
    
    print(f"\n📊 Total recipes extracted: {len(all_recipes)}")
    
    # Group recipes by normalized title
    recipe_groups = defaultdict(list)
    for recipe in all_recipes:
        normalized = normalize_title(recipe['title'])
        recipe_groups[normalized].append(recipe)
    
    print(f"📊 Unique recipe titles: {len(recipe_groups)}")
    
    # Count duplicates and versions
    duplicates_found = 0
    versioned_recipes = 0
    
    for normalized_title, recipes in recipe_groups.items():
        if len(recipes) > 1:
            content_hashes = set()
            for recipe in recipes:
                content_hashes.add(get_content_hash(recipe['content']))
            
            if len(content_hashes) == 1:
                duplicates_found += len(recipes) - 1
            else:
                versioned_recipes += 1
    
    # Create combined document with unified format
    combined_doc = Document()
    
    # Add title page
    title = combined_doc.add_heading("THE LARIAT RECIPE BOOK", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = combined_doc.add_paragraph("Unified Professional Edition")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = combined_doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    stats = combined_doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats.add_run(f"\n{len(recipe_groups)} Recipes • {duplicates_found} Duplicates Removed • {versioned_recipes} Versioned Recipes")
    
    combined_doc.add_page_break()
    
    # Add table of contents
    combined_doc.add_heading("Table of Contents", 1)
    
    sorted_titles = sorted(recipe_groups.keys())
    
    for normalized_title in sorted_titles:
        recipes = recipe_groups[normalized_title]
        display_title = recipes[0]['title']
        
        # Check for multiple versions
        content_hashes = set()
        for recipe in recipes:
            content_hashes.add(get_content_hash(recipe['content']))
        
        if len(content_hashes) > 1:
            combined_doc.add_paragraph(f"• {display_title} ({len(content_hashes)} versions)")
        else:
            combined_doc.add_paragraph(f"• {display_title}")
    
    combined_doc.add_page_break()
    
    # Add recipes in unified format
    combined_doc.add_heading("Recipes", 1)
    
    for normalized_title in sorted_titles:
        recipes = recipe_groups[normalized_title]
        
        # Check for content differences
        content_map = {}
        for recipe in recipes:
            hash_val = get_content_hash(recipe['content'])
            if hash_val not in content_map:
                content_map[hash_val] = []
            content_map[hash_val].append(recipe)
        
        if len(content_map) == 1:
            # Single unique version - use the first one
            recipe = recipes[0]
            add_unified_recipe(combined_doc, recipe, recipes)
            
        else:
            # Multiple different versions
            version_num = 1
            for hash_val, hash_recipes in content_map.items():
                recipe = hash_recipes[0]
                
                version_title = f"{recipe['title']} (Version {version_num})"
                combined_doc.add_heading(version_title, 2)
                
                # Note which sources have this version
                sources = list(set(r['source'] for r in hash_recipes))
                source_para = combined_doc.add_paragraph()
                source_run = source_para.add_run(f"Found in: {', '.join(sources)}")
                source_run.italic = True
                source_run.font.size = Pt(9)
                
                add_unified_recipe_content(combined_doc, recipe['content'])
                
                version_num += 1
        
        # Add separator
        sep = combined_doc.add_paragraph()
        sep_run = sep.add_run("─" * 40)
        sep_run.font.size = Pt(8)
    
    # Save the combined document
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    combined_doc.save(OUTPUT_FILE)
    
    print("\n" + "=" * 60)
    print("✅ COMBINATION COMPLETE")
    print("=" * 60)
    print(f"📄 Output: {OUTPUT_FILE}")
    print(f"📊 Unique recipes: {len(recipe_groups)}")
    print(f"🔄 Exact duplicates removed: {duplicates_found}")
    print(f"📝 Recipes with multiple versions: {versioned_recipes}")
    print("=" * 60)
    
    # Print recipe list
    print("\n📋 RECIPE LIST:")
    for i, title in enumerate(sorted_titles, 1):
        display = recipe_groups[title][0]['title']
        print(f"  {i:2}. {display}")
    
    return OUTPUT_FILE


def add_unified_recipe(combined_doc, recipe, all_versions):
    """Add a recipe in unified format with total yields and instructions"""
    combined_doc.add_heading(recipe['title'], 2)
    
    # Add source note
    sources = list(set(r['source'] for r in all_versions))
    source_para = combined_doc.add_paragraph()
    source_run = source_para.add_run(f"[Source: {', '.join(sources)}]")
    source_run.italic = True
    source_run.font.size = Pt(8)
    
    # Process content into unified format
    add_unified_recipe_content(combined_doc, recipe['content'])


def add_unified_recipe_content(combined_doc, content_lines):
    """Process recipe content into unified format with sections"""
    # Categorize content lines
    yield_info = []
    ingredients = []
    instructions = []
    notes = []
    
    current_section = None
    
    for line in content_lines:
        line_lower = line.lower().strip()
        
        # Detect section headers
        if any(word in line_lower for word in ['yield', 'serves', 'makes', 'yields']):
            current_section = 'yield'
            yield_info.append(line)
        elif any(word in line_lower for word in ['ingredients', 'base', 'components']):
            current_section = 'ingredients'
            if 'ingredients' not in line_lower:
                ingredients.append(line)  # Add the section header
        elif any(word in line_lower for word in ['directions', 'procedure', 'method', 'instructions', 'steps']):
            current_section = 'instructions'
            if not any(word in line_lower for word in ['directions', 'procedure', 'method', 'instructions', 'steps']):
                instructions.append(line)  # Add the section header
        elif any(word in line_lower for word in ['notes', 'tips', 'garnish']):
            current_section = 'notes'
            notes.append(line)
        else:
            # Add to current section
            if current_section == 'yield':
                yield_info.append(line)
            elif current_section == 'ingredients':
                ingredients.append(line)
            elif current_section == 'instructions':
                instructions.append(line)
            elif current_section == 'notes':
                notes.append(line)
            else:
                # Try to auto-detect based on content
                if re.match(r'^[\d½¼¾⅓⅔⅛]', line.strip()):  # Starts with quantity
                    ingredients.append(line)
                elif any(line_lower.startswith(word) for word in ['add', 'mix', 'stir', 'cook', 'heat', 'bring', 'simmer', 'reduce', 'season', 'adjust']):
                    instructions.append(line)
                else:
                    # Default to ingredients if we can't determine
                    if not ingredients and not instructions:
                        ingredients.append(line)
                    elif ingredients and not instructions:
                        ingredients.append(line)
                    else:
                        instructions.append(line)
    
    # Add sections in standard order
    
    # Yield/Portion Info
    if yield_info:
        yield_heading = combined_doc.add_paragraph()
        yield_heading.add_run("YIELD").bold = True
        for line in yield_info:
            combined_doc.add_paragraph(line)
    
    # Ingredients
    if ingredients:
        ingredients_heading = combined_doc.add_paragraph()
        ingredients_heading.add_run("INGREDIENTS").bold = True
        for line in ingredients:
            combined_doc.add_paragraph(line)
    
    # Instructions
    if instructions:
        instructions_heading = combined_doc.add_paragraph()
        instructions_heading.add_run("INSTRUCTIONS").bold = True
        for line in instructions:
            combined_doc.add_paragraph(line)
    
    # Notes
    if notes:
        notes_heading = combined_doc.add_paragraph()
        notes_heading.add_run("NOTES").bold = True
        for line in notes:
            combined_doc.add_paragraph(line)


if __name__ == "__main__":
    combine_recipes()
