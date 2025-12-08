#!/usr/bin/env python3
"""
Extract text from Lariat Recipe Book.docx
"""

from docx import Document
import json

def extract_recipe_book(docx_path):
    """Extract all text from the recipe book"""
    doc = Document(docx_path)

    full_text = []
    recipes = []
    current_recipe = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text.append(text)

        # Try to detect recipe structure
        if para.style.name.startswith('Heading'):
            if current_recipe:
                recipes.append(current_recipe)
            current_recipe = {
                'title': text,
                'content': []
            }
        elif current_recipe:
            current_recipe['content'].append(text)

    if current_recipe:
        recipes.append(current_recipe)

    return {
        'full_text': full_text,
        'recipes': recipes,
        'paragraph_count': len(doc.paragraphs)
    }

if __name__ == "__main__":
    docx_path = "/Users/seanburdges/Desktop/LARIAT/Lariat Recipe Book.docx"

    print("=" * 80)
    print("EXTRACTING LARIAT RECIPE BOOK")
    print("=" * 80)

    data = extract_recipe_book(docx_path)

    print(f"\nTotal paragraphs: {data['paragraph_count']}")
    print(f"Recipes found: {len(data['recipes'])}")

    # Save full text
    with open("recipe_book_full_text.txt", "w") as f:
        f.write("\n".join(data['full_text']))

    # Save structured recipes
    with open("recipe_book_recipes.json", "w") as f:
        json.dump(data['recipes'], f, indent=2)

    print("\n✓ Saved recipe_book_full_text.txt")
    print("✓ Saved recipe_book_recipes.json")

    # Print first few recipes
    print(f"\nFirst {min(5, len(data['recipes']))} recipes:")
    for i, recipe in enumerate(data['recipes'][:5], 1):
        print(f"\n{i}. {recipe['title']}")
        print(f"   Lines: {len(recipe['content'])}")
        if recipe['content']:
            print(f"   Preview: {recipe['content'][0][:80]}...")
